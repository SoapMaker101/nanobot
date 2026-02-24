"""File system tools: read, write, edit."""

import difflib
import mimetypes
from pathlib import Path
from typing import Any

from nanobot.agent.tools.base import Tool


def _read_pdf(file_path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        import pdfplumber
    except ImportError:
        return "Error: Reading PDF requires the pdfplumber package. Install with: pip install pdfplumber"
    try:
        with pdfplumber.open(file_path) as pdf:
            parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n\n".join(parts) if parts else "The PDF file contains no extractable text (may be scanned/image-only)."
    except Exception as e:
        return f"Error reading PDF: {e}"


def _read_docx(file_path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        import docx
    except ImportError:
        return "Error: Reading DOCX requires the python-docx package. Install with: pip install python-docx"
    try:
        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return "The DOCX file contains no text."
        return "\n".join(paragraphs)
    except Exception as e:
        return f"Error reading DOCX: {e}"


def _read_excel(file_path: Path) -> str:
    """Extract data from an Excel file (.xlsx / .xls)."""
    try:
        import openpyxl
    except ImportError:
        return "Error: Reading Excel requires the openpyxl package. Install with: pip install openpyxl"
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if rows:
                parts.append(f"Sheet: {sheet.title}")
                for row in rows:
                    parts.append("\t".join(str(c) if c is not None else "" for c in row))
        wb.close()
        return "\n".join(parts) if parts else "The Excel file contains no data."
    except Exception as e:
        return f"Error reading Excel: {e}"


def _read_text_with_encoding(file_path: Path) -> str:
    """Read a text file, trying UTF-8 then common fallbacks."""
    raw = file_path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _smart_read(file_path: Path) -> str:
    """Read file content based on extension (PDF, DOCX, XLSX, or plain text)."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return _read_pdf(file_path)
    if ext in (".docx", ".doc"):
        return _read_docx(file_path)
    if ext in (".xlsx", ".xls"):
        return _read_excel(file_path)
    if ext and mimetypes.guess_type(str(file_path))[0] and (mimetypes.guess_type(str(file_path))[0] or "").startswith("image/"):
        return f"This is an image file ({file_path.suffix}). Use other tools to process images."
    return _read_text_with_encoding(file_path)


def _write_docx(file_path: Path, content: str) -> str:
    """Create a DOCX file with the given text content."""
    try:
        import docx
    except ImportError:
        return "Error: Creating DOCX requires python-docx. Install with: pip install python-docx"
    doc = docx.Document()
    for block in content.strip().split("\n\n"):
        doc.add_paragraph(block.replace("\n", " "))
    doc.save(str(file_path))
    return f"Successfully wrote DOCX to {file_path}"


def _write_xlsx(file_path: Path, content: str) -> str:
    """Create an XLSX file. Content: one value per line (column A), or tab-separated for multiple columns."""
    try:
        import openpyxl
    except ImportError:
        return "Error: Creating Excel requires openpyxl. Install with: pip install openpyxl"
    wb = openpyxl.Workbook()
    ws = wb.active
    for row_idx, line in enumerate(content.strip().splitlines(), start=1):
        cells = line.split("\t")
        for col_idx, val in enumerate(cells, start=1):
            ws.cell(row=row_idx, column=col_idx, value=val.strip() if isinstance(val, str) else val)
    wb.save(str(file_path))
    return f"Successfully wrote Excel to {file_path}"


def _resolve_path(path: str, workspace: Path | None = None, allowed_dir: Path | None = None) -> Path:
    """Resolve path against workspace (if relative) and enforce directory restriction."""
    p = Path(path).expanduser()
    if not p.is_absolute() and workspace:
        p = workspace / p
    resolved = p.resolve()
    if allowed_dir and not str(resolved).startswith(str(allowed_dir.resolve())):
        raise PermissionError(f"Path {path} is outside allowed directory {allowed_dir}")
    return resolved


class ReadFileTool(Tool):
    """Tool to read file contents."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "read_file"
    
    @property
    def description(self) -> str:
        return (
            "Read the contents of a file. When the user sends you a file, the message contains a path like [file: /path/to/file]. "
            "Call read_file with that path to get the text. Supports .txt, .docx, .pdf, .xlsx (and .doc, .xls)."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Full path to the file (e.g. from [file: /path] in the user message). Use as-is including extension."
                }
            },
            "required": ["path"]
        }
    
    async def execute(self, path: str, **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.is_file():
                return f"Error: Not a file: {path}"
            return _smart_read(file_path)
        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error reading file: {str(e)}"


class WriteFileTool(Tool):
    """Tool to write content to a file."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "write_file"
    
    @property
    def description(self) -> str:
        return (
            "CREATE a new file: pass path (e.g. hello.txt, report.docx) and content. Use this when the user asks to create or send a file — create with write_file first, then send with message(media=[path]). "
            "Supports .txt, .docx, .xlsx. Do NOT use edit_file to create files; edit_file only replaces text in existing files."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The file path to write to (e.g. report.txt, report.docx, data.xlsx)"
                },
                "content": {
                    "type": "string",
                    "description": "The content to write (text; for xlsx, use one value per line for first column, or tab-separated for columns)"
                }
            },
            "required": ["path", "content"]
        }

    async def execute(self, path: str, content: str, **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            file_path.parent.mkdir(parents=True, exist_ok=True)
            ext = file_path.suffix.lower()
            if ext == ".docx":
                return _write_docx(file_path, content)
            if ext == ".xlsx":
                return _write_xlsx(file_path, content)
            file_path.write_text(content, encoding="utf-8")
            return f"Successfully wrote {len(content)} bytes to {file_path}"
        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error writing file: {str(e)}"


class EditFileTool(Tool):
    """Tool to edit a file by replacing text."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "edit_file"
    
    @property
    def description(self) -> str:
        return (
            "Replace old_text with new_text in an EXISTING file. Requires path, old_text, and new_text. "
            "Use only when the file already exists. To CREATE a new file with content, use write_file instead."
        )
    
    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The file path to edit"
                },
                "old_text": {
                    "type": "string",
                    "description": "The exact text to find and replace"
                },
                "new_text": {
                    "type": "string",
                    "description": "The text to replace with"
                }
            },
            "required": ["path", "old_text", "new_text"]
        }
    
    async def execute(self, path: str, old_text: str, new_text: str, **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"

            content = file_path.read_text(encoding="utf-8")

            if old_text not in content:
                return self._not_found_message(old_text, content, path)

            # Count occurrences
            count = content.count(old_text)
            if count > 1:
                return f"Warning: old_text appears {count} times. Please provide more context to make it unique."

            new_content = content.replace(old_text, new_text, 1)
            file_path.write_text(new_content, encoding="utf-8")

            return f"Successfully edited {file_path}"
        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error editing file: {str(e)}"

    @staticmethod
    def _not_found_message(old_text: str, content: str, path: str) -> str:
        """Build a helpful error when old_text is not found."""
        lines = content.splitlines(keepends=True)
        old_lines = old_text.splitlines(keepends=True)
        window = len(old_lines)

        best_ratio, best_start = 0.0, 0
        for i in range(max(1, len(lines) - window + 1)):
            ratio = difflib.SequenceMatcher(None, old_lines, lines[i : i + window]).ratio()
            if ratio > best_ratio:
                best_ratio, best_start = ratio, i

        if best_ratio > 0.5:
            diff = "\n".join(difflib.unified_diff(
                old_lines, lines[best_start : best_start + window],
                fromfile="old_text (provided)", tofile=f"{path} (actual, line {best_start + 1})",
                lineterm="",
            ))
            return f"Error: old_text not found in {path}.\nBest match ({best_ratio:.0%} similar) at line {best_start + 1}:\n{diff}"
        return f"Error: old_text not found in {path}. No similar text found. Verify the file content."


class ListDirTool(Tool):
    """Tool to list directory contents."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "list_dir"
    
    @property
    def description(self) -> str:
        return "List the contents of a directory."
    
    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The directory path to list"
                }
            },
            "required": ["path"]
        }
    
    async def execute(self, path: str, **kwargs: Any) -> str:
        try:
            dir_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not dir_path.exists():
                return f"Error: Directory not found: {path}"
            if not dir_path.is_dir():
                return f"Error: Not a directory: {path}"

            items = []
            for item in sorted(dir_path.iterdir()):
                prefix = "📁 " if item.is_dir() else "📄 "
                items.append(f"{prefix}{item.name}")

            if not items:
                return f"Directory {path} is empty"

            return "\n".join(items)
        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error listing directory: {str(e)}"
